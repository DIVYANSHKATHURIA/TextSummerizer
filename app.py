from flask import Flask, render_template, request
import requests
from werkzeug.utils import secure_filename
import os
from dotenv import load_dotenv
import PyPDF2
import docx

load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'

# Create the upload directory if it doesn't exist
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

API_URL = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"
HEADERS = {"Authorization": f"Bearer {os.getenv('HUGGINGFACE_API_KEY')}"}

@app.route("/", methods=['GET', 'POST'])
def index():
    return render_template('index.html')

@app.route("/summerize", methods=["POST"])
def summerize():
    input_word_count = 0
    result = ''
    word_count = 0
    slider_value = request.form.get('summary_length', '0')  # Default to '0' if not set

    # Check if a file was uploaded
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text based on file type
        if filename.endswith('.pdf'):
            texts = extract_text_from_pdf(filepath)
        elif filename.endswith('.docx'):
            texts = extract_text_from_docx(filepath)
        else:
            return render_template('index.html', result="Unsupported file type")

        if isinstance(texts, str):  # An error message was returned
            return render_template('index.html', result=texts)

        # Concatenate all extracted text into one string
        all_text = "\n".join(texts)
        input_word_count = len(all_text.split())

        summaries = []
        for text in texts:
            min_length, max_length = get_summary_length_parameters(len(text.split()), slider_value)
            if text.strip():  # Skip empty text
                summary = query({
                    "inputs": text,
                    "parameters": {
                        "max_length": max_length,
                        "min_length": min_length
                    }
                })
                if isinstance(summary, list) and summary:
                    summaries.append(summary[0].get('summary_text', 'Error occurred'))
                else:
                    summaries.append('Error occurred')
        
        result = "\n\n".join(summaries)
        word_count = len(result.split())

    elif 'data' in request.form and request.form['data']:
        data = request.form['data']
        input_word_count = len(data.split())

        # Determine length parameters based on slider value
        min_length, max_length = get_summary_length_parameters(input_word_count, slider_value)

        summary = query({
            "inputs": data,
            "parameters": {
                "max_length": max_length,
                "min_length": min_length
            }
        })
        if isinstance(summary, list) and summary:
            result = summary[0].get('summary_text', 'Error occurred')
        else:
            result = 'Error occurred'

        word_count = len(result.split())

    else:
        return render_template('index.html', result="No file selected or no text entered")

    return render_template('index.html', result=result, word_count=word_count, input_word_count=input_word_count)

def extract_text_from_pdf(pdf_path):
    texts = []
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            if num_pages > 10:
                return "Error: PDF exceeds the maximum page limit of 10 pages."
            for page_number in range(num_pages):
                page = reader.pages[page_number]
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
                else:
                    texts.append("")  # Handle pages with no text
    except Exception as e:
        return f"Error extracting text from PDF: {e}"
    return texts

def extract_text_from_docx(docx_path):
    texts = []
    try:
        doc = docx.Document(docx_path)
        num_paragraphs = len(doc.paragraphs)
        if num_paragraphs > 100:
            return "Error: DOCX exceeds the maximum paragraph limit of 100 paragraphs."
        for paragraph in doc.paragraphs:
            texts.append(paragraph.text + '\n')
    except Exception as e:
        return f"Error extracting text from DOCX: {e}"
    return texts

def query(payload):
    try:
        response = requests.post(API_URL, headers=HEADERS, json=payload)
        response.raise_for_status()  # Raise an error for bad status codes
        return response.json()  # This will be a list or dictionary based on the API response
    except requests.exceptions.RequestException as e:
        return [{"summary_text": f"Error: {e}"}]  # Return a list with an error message if an exception occurs

def get_summary_length_parameters(word_count, slider_value):
    """Determine min and max length for summary based on slider value."""
    if slider_value == '0':  # Short
        min_length = word_count // 3
        max_length = min_length + 50
    elif slider_value == '1':  # Medium
        min_length = (2 * word_count) // 5
        max_length = min_length + 50
    elif slider_value == '2':  # Long
        min_length = word_count // 2
        max_length = min_length + 50
    else:
        min_length = 100
        max_length = 150

    return min_length, max_length

if __name__ == '__main__':
    # app.run(debug=True)
    pass
